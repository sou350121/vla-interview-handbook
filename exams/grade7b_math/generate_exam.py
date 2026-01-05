from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from assets_builder import ensure_assets


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def _as_letters(n: int) -> list[str]:
    return [chr(ord("A") + i) for i in range(n)]


def _escape_xml(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


@dataclass(frozen=True)
class ImageRef:
    file: str
    caption: str | None = None
    width_in: float | None = None


def _iter_questions_in_order(bank: dict[str, Any]) -> Iterable[dict[str, Any]]:
    by_id = {q["id"]: q for q in bank["questions"]}
    for sec in bank["sections"]:
        for qid in sec["question_ids"]:
            yield by_id[qid]


def build_docx(bank: dict[str, Any], out_path: Path, base_dir: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(bank["metadata"]["title"])
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(bank["metadata"]["subtitle"])

    doc.add_paragraph(bank["metadata"].get("note", ""))

    doc.add_paragraph("注意事项：")
    doc.add_paragraph("1. 请在答题纸/空白处按题号作答，步骤写清楚。")
    doc.add_paragraph("2. 选择题将正确选项填涂/写在题号后。")

    doc.add_paragraph()

    # Questions
    by_id = {q["id"]: q for q in bank["questions"]}
    q_number = 1
    for section in bank["sections"]:
        doc.add_heading(section["title"], level=2)
        for qid in section["question_ids"]:
            q = by_id[qid]
            p = doc.add_paragraph(f"{q_number}. {q['stem']}")

            # images (if any)
            for img in q.get("images", []):
                ref = ImageRef(**img)
                img_path = base_dir / ref.file
                if img_path.exists():
                    width = Inches(ref.width_in) if ref.width_in else None
                    doc.add_picture(str(img_path), width=width)
                    if ref.caption:
                        cap = doc.add_paragraph(ref.caption)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if q["type"] == "choice":
                letters = _as_letters(len(q["choices"]))
                for letter, choice in zip(letters, q["choices"]):
                    doc.add_paragraph(f"   {letter}. {choice}")
            q_number += 1

        doc.add_paragraph()

    # Answers & analysis
    doc.add_page_break()
    doc.add_heading("参考答案与解析", level=1)

    q_number = 1
    for q in _iter_questions_in_order(bank):
        doc.add_paragraph(f"{q_number}. 【答案】{q['answer']}")
        doc.add_paragraph(f"   【解析】{q['analysis']}")
        if q.get("analogy"):
            doc.add_paragraph(f"   【生活类比】{q['analogy']}")
        doc.add_paragraph()
        q_number += 1

    doc.save(out_path)


def build_pdf(bank: dict[str, Any], out_path: Path, base_dir: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.lib.enums import TA_CENTER

    # Chinese font without external files (CID font).
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "CnNormal",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=11,
        leading=15,
    )
    title = ParagraphStyle(
        "CnTitle",
        parent=normal,
        fontSize=16,
        leading=20,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    subtitle = ParagraphStyle(
        "CnSubtitle",
        parent=normal,
        fontSize=12,
        leading=16,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h2 = ParagraphStyle(
        "CnH2",
        parent=normal,
        fontSize=13,
        leading=18,
        spaceBefore=10,
        spaceAfter=6,
    )

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
    )

    story: list[Any] = []
    story.append(Paragraph(_escape_xml(bank["metadata"]["title"]), title))
    story.append(Paragraph(_escape_xml(bank["metadata"]["subtitle"]), subtitle))
    if bank["metadata"].get("note"):
        story.append(Paragraph(_escape_xml(bank["metadata"]["note"]), normal))
        story.append(Spacer(1, 10))

    story.append(Paragraph("注意事项：", normal))
    story.append(Paragraph("1. 请在答题处按题号作答，步骤写清楚。", normal))
    story.append(Paragraph("2. 选择题将正确选项写在题号后。", normal))
    story.append(Spacer(1, 14))

    by_id = {q["id"]: q for q in bank["questions"]}
    q_number = 1
    for section in bank["sections"]:
        story.append(Paragraph(_escape_xml(section["title"]), h2))
        for qid in section["question_ids"]:
            q = by_id[qid]
            stem = _escape_xml(q["stem"]).replace("\n", "<br/>")
            story.append(Paragraph(f"{q_number}. {stem}", normal))

            for img in q.get("images", []):
                ref = ImageRef(**img)
                img_path = base_dir / ref.file
                if img_path.exists():
                    from PIL import Image as PILImage

                    # width in inches -> points
                    w_in = ref.width_in if ref.width_in else 4.6
                    w_pt = w_in * inch
                    iw, ih = PILImage.open(img_path).size
                    h_pt = ih * (w_pt / iw)
                    story.append(Spacer(1, 6))
                    story.append(RLImage(str(img_path), width=w_pt, height=h_pt))
                    if ref.caption:
                        story.append(Paragraph(_escape_xml(ref.caption), ParagraphStyle("Cap", parent=normal, alignment=TA_CENTER)))
                    story.append(Spacer(1, 6))

            if q["type"] == "choice":
                letters = _as_letters(len(q["choices"]))
                opts = "<br/>".join(
                    f"{letter}. {_escape_xml(choice)}" for letter, choice in zip(letters, q["choices"])
                )
                story.append(Paragraph(opts, normal))
            story.append(Spacer(1, 10))
            q_number += 1

    # Answers
    story.append(Spacer(1, 18))
    story.append(Paragraph("参考答案与解析", h2))

    q_number = 1
    for q in _iter_questions_in_order(bank):
        ans = _escape_xml(str(q["answer"])).replace("\n", "<br/>")
        ana = _escape_xml(str(q["analysis"])).replace("\n", "<br/>")
        story.append(Paragraph(f"{q_number}. 【答案】{ans}", normal))
        story.append(Paragraph(f"【解析】{ana}", normal))
        if q.get("analogy"):
            alg = _escape_xml(str(q["analogy"])).replace("\n", "<br/>")
            story.append(Paragraph(f"【生活类比】{alg}", normal))
        story.append(Spacer(1, 10))
        q_number += 1

    doc.build(story)


def main() -> None:
    base_dir = Path(__file__).resolve().parent
    assets_dir = base_dir / "assets"
    _ensure_dir(assets_dir)
    ensure_assets(assets_dir, overwrite=True)

    bank = _read_json(base_dir / "questions.json")

    def choose_basename(stem: str = "exam") -> str:
        """
        If exam.docx is currently opened/locked (common on Windows), avoid overwriting it.
        Pick the first writable name: exam.docx, exam_v2.docx, exam_v3.docx, ...
        """
        candidates = [stem] + [f"{stem}_v{i}" for i in range(2, 50)]
        for name in candidates:
            p = base_dir / f"{name}.docx"
            try:
                # Try opening for write (create if missing). This will fail if locked.
                with open(p, "ab"):
                    pass
                return name
            except PermissionError:
                continue
        # last resort
        return f"{stem}_v99"

    basename = choose_basename("exam")
    out_docx = base_dir / f"{basename}.docx"
    out_pdf = base_dir / f"{basename}.pdf"

    build_docx(bank, out_docx, base_dir)
    build_pdf(bank, out_pdf, base_dir)

    print(f"OK: {out_docx}")
    print(f"OK: {out_pdf}")


if __name__ == "__main__":
    main()


